# -*- coding: utf-8 -*-
import os
import pandas as pd
from docx import Document
from docx.shared import Inches

# 设置包含多个 CSV 的文件夹路径
csv_folder = r"标准映射后"  # 替换为你的目录路径
output_folder = r"word"  # 输出 Word 文件目录

if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# 遍历文件夹中的所有 CSV 文件
for file in os.listdir(csv_folder):
    if file.lower().endswith(".csv"):
        csv_path = os.path.join(csv_folder, file)

        # 读取 CSV（假设为 GBK 编码）
        try:
            df = pd.read_csv(csv_path, encoding="gbk")
        except Exception as e:
            print(u"⚠️ 读取失败：{}，错误：{}".format(file, e))
            continue

        # 创建 Word 文档
        doc = Document()
        # 按Filename分组处理
        grouped = df.groupby(["Filename", "Filename_zh"])  # 同时按中文名分组

        for (filename, filename_zh), group in grouped:
            # 添加双层标题（英文名+中文名）
            doc.add_heading(f"{filename}", level=2)  # 英文图层名
            doc.add_heading(f"{filename_zh}", level=3)  # 新增中文图层名

            # 创建 5 列的表格（不含序号）
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # 设置表头（去掉“序号”）
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u"字段名称"
            hdr_cells[1].text = u"字段代码"
            hdr_cells[2].text = u"字段类型"
            hdr_cells[3].text = u"字段长度"
            hdr_cells[4].text = u"小数位数"

            # 添加字段行（去掉 i）
            for row in group.itertuples():
                cells = table.add_row().cells
                cells[0].text = str(row.FieldName)
                cells[1].text = str(row.FieldAlias)
                cells[2].text = str(row.Type)
                cells[3].text = str(row.Length)
                cells[4].text = str(row.Scale)

            # 添加段落分隔
            doc.add_paragraph()

        # 保存 Word 文档
        output_docx = os.path.join(output_folder, os.path.splitext(file)[0] + ".docx")
        doc.save(output_docx)
        print(u"✅ 已生成 Word 文件：{}".format(output_docx))
