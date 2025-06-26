
import pandas as pd
from docx import Document
from docx.shared import Inches
from collections import defaultdict

# 读取 CSV 文件
csv_file = r"C:\Users\LMQ\Desktop\【整体评价】天津市2016年度区域更新评价_空间数据.csv"  # 替换为你的路径
df = pd.read_csv(csv_file, encoding="gbk")  # 如果你保存是GBK编码

# 创建 Word 文档
doc = Document()

# 按照 Filename 分组
grouped = df.groupby("Filename")

for filename, group in grouped:
    # 添加标题
    doc.add_heading(u"{}.shp".format(filename.strip(".shp")), level=2)

    # 创建 5 列的表格（不含序号）
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # 设置表头（去掉“序号”）
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u"字段名称"
    hdr_cells[1].text = u"字段代码"
    hdr_cells[2].text = u"类型"
    hdr_cells[3].text = u"长度"
    hdr_cells[4].text = u"小数"

    # 添加字段行（去掉 i）
    for row in group.itertuples():
        cells = table.add_row().cells
        cells[0].text = str(row.FieldName)
        cells[1].text = str(row.FieldAlias)
        cells[2].text = str(row.Type)
        cells[3].text = str(row.Length)
        cells[4].text = str(row.Scale)

    # 添加段落分隔
    doc.add_paragraph()


# 保存 Word 文档
output_path = "字段信息表.docx"
doc.save(output_path)
print(u"✅ Word 文件已生成：{}".format(output_path))
